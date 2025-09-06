from docx import Document
from docx.shared import Pt

# Create a Word document
doc = Document()
doc.add_heading("Movie Helper Flowchart Description", 0)

# Start
doc.add_heading("Start:", level=1)
doc.add_paragraph(
    "Oval → Label: 'User Input: Movie Helper'\n"
    "This is where the user starts interacting with the system."
)

# Decision Points
doc.add_heading("Decision Points (Diamonds):", level=1)

decision_points = [
    "Do you want Action or Non-Action? (Yes → Action / No → Non-Action)",
    "Do you prefer Happy or Sad movies? (Yes → Happy / No → Sad)",
    "Do you want a Movie for a group larger than 3? (Yes → Large Group / No → Small Group)",
    "Do you want the latest releases? (Yes → Latest / No → Older Movies)",
    "Do you prefer movies with high ratings? (Yes → Highly Rated / No → Any Rating)",
    "Do you want a Movie available on Streaming? (Yes → Online / No → Theater Only)",
    "Do you prefer long movies (>2 hours)? (Yes → Long / No → Short)",
    "Do you prefer English language? (Yes → English / No → Local language)"
]

for idx, q in enumerate(decision_points, 1):
    doc.add_paragraph(f"{idx}. {q}")

# Process / Action Boxes
doc.add_heading("Process / Action Boxes (Rectangles):", level=1)
doc.add_paragraph(
    "Collect all the inputs from the user and filter available movies.\n"
    "Apply all criteria based on yes/no answers."
)

# Recommendations
doc.add_heading("Recommendations (Rectangles):", level=1)
recommendations = [
    "Action, Happy, Large Group, Latest, High Rating → Example Movie: Avengers: Endgame",
    "Non-Action, Sad, Small Group, Any Release, Streaming → Example Movie: The Pursuit of Happyness",
    "Action, Sad, Large Group, Highly Rated → Example Movie: Inception",
    "Non-Action, Happy, Any Group, Any Rating, Streaming → Example Movie: The Intern"
]

for rec in recommendations:
    doc.add_paragraph(rec)

# End
doc.add_heading("End:", level=1)
doc.add_paragraph("Oval → Label: 'End of Recommendation'")

# Tips
doc.add_heading("Tips for Diagram:", level=1)
doc.add_paragraph(
    "- Use Oval for Start/End\n"
    "- Use Diamond for Questions / Decisions\n"
    "- Use Rectangle for Actions / Recommendations\n"
    "- Connect all with arrows showing yes/no flow"
)

# Save Word file in Downloads
import os
downloads_path = os.path.join(os.path.expanduser("~"), "Downloads", "Movie_Helper_Flowchart5.docx")
doc.save(downloads_path)

print(f"Word document created in Downloads folder: {downloads_path}")
